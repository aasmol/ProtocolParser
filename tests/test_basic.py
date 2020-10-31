# -*- coding: utf-8 -*-

from .context import dataextractor

import docx


class TestFunctions(object):

    def test_get_name(self):
        doc = docx.Document(docx='c:/Users/user/Documents/Word parsing/Test docs/1608010001-ПРЧ.docx')
        assert dataextractor.edata.get_device_name(doc) == 'A54SX16A-TQ100I'

    def test_find_table(self):
        doc = docx.Document(docx='c:/Users/user/Documents/Word parsing/Test docs/1608010001-ПРЧ.docx')
        table = dataextractor.edata.get_test_table(doc, 'ТЭ')
        row_ind = 0
        assert table.cell(row_ind, 0).text == 'Ион'

    def test_table_to_df(self):
        doc = docx.Document(docx='c:/Users/user/Documents/Word parsing/Test docs/1641010002-ПРЧ.docx')
        table = dataextractor.edata.get_test_table(doc, 'ТЭ')
        table_data = dataextractor.edata.table_to_df(table)
        print(table_data)
        print(table_data.dtypes)
        assert table.cell(0,0).text == list(table_data)[0]

    def test_prepare_data(self):
        doc = docx.Document(docx='c:/Users/user/Documents/Word parsing/Test docs/1641010002-ПРЧ.docx')
        table = dataextractor.edata.get_test_table(doc, 'ТЭ')
        processed_data = dataextractor.edata.prepare_data(table)
        print(processed_data)
        assert True





